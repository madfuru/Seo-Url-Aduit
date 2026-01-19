from docx import Document

def fill_docx(template_path: str, out_path: str, replacements: dict):
    doc = Document(template_path)

    def repl(s: str) -> str:
        for k, v in replacements.items():
            s = s.replace(f"{{{{{k}}}}}", str(v))
        return s

    for p in doc.paragraphs:
        if "{{" in p.text:
            txt = repl(p.text)
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = txt
            else:
                p.add_run(txt)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "{{" in cell.text:
                    cell.text = repl(cell.text)

    doc.save(out_path)
